from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# ── Configuration ───────────────────────────────────────
IMAGES_DIR = r"E:\Projects\facedetect\images"
OUTPUT_DOCX = r"E:\Projects\facedetect\Face_Mask_Detection_Report.docx"

doc = Document()

# ── 1. Title Page ───────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Face Mask Compliance Detection Using LBP + HOG + SVM")
run.bold = True
run.font.size = Pt(24)

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author.add_run("\nRegister No: 2116230701158\nDepartment of CSE\nRajalakshmi Engineering College")
author_run.font.size = Pt(14)

doc.add_page_break()

# ── 2. Introduction ─────────────────────────────────────
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    "Face mask compliance monitoring is essential for public safety and health governance. "
    "This project implements a lightweight classical computer vision system using Local Binary Patterns (LBP), "
    "Histogram of Oriented Gradients (HOG), and a Support Vector Machine (SVM) for high-precision, real-time detection."
)

# ── 3. Dataset Description ──────────────────────────────
doc.add_heading('2. Dataset Description', level=1)
doc.add_paragraph("The dataset is well-balanced with thousands of high-quality samples.")

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Class'
hdr_cells[1].text = 'Images'
hdr_cells[2].text = 'Description'
for text in ['with_mask', '3725', 'Mask worn']:
    row_cells = table.add_row().cells
    row_cells[0].text = 'with_mask'
    row_cells[1].text = '3725'
    row_cells[2].text = 'Mask worn correctly'
    break # Just one row for demo
row_cells = table.add_row().cells
row_cells[0].text = 'without_mask'
row_cells[1].text = '3828'
row_cells[2].text = 'No mask detected'

img_path = os.path.join(IMAGES_DIR, "dataset_bar.jpg")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 4. Preprocessing ────────────────────────────────────
doc.add_heading('3. Preprocessing', level=1)
doc.add_paragraph("Steps applied: Resize (64x64), Grayscale, and CLAHE enhancement.")

img_path = os.path.join(IMAGES_DIR, "clahe.jpg")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 5. Feature Extraction ──────────────────────────────
doc.add_heading('4. Feature Extraction', level=1)
doc.add_paragraph("Texture (LBP) and Structural (HOG) features are combined for robust detection.")

doc.add_heading('4.1 LBP Visualization', level=2)
img_path = os.path.join(IMAGES_DIR, "lbp.jpg")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('4.2 HOG Visualization', level=2)
img_path = os.path.join(IMAGES_DIR, "hog.jpg")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 6. Results ──────────────────────────────────────────
doc.add_heading('5. Results and Evaluation', level=1)
doc.add_paragraph("The model achieved a final accuracy of 90.21%.")

img_path = os.path.join(IMAGES_DIR, "confusion.jpg")
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 7. Core Conclusion ──────────────────────────────────
doc.add_heading('6. Conclusion', level=1)
doc.add_paragraph(
    "The 90.21% accuracy shows that classical feature extraction (LBP+HOG) "
    "remains a powerful tool for visual classification tasks with low CPU overhead."
)

# ── 8. Real-time Detection ─────────────────────────────
doc.add_heading('7. Real-time Detection Demos', level=1)
for img_name in ["mask.jpg", "nomask.jpg"]:
    img_path = os.path.join(IMAGES_DIR, img_name)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUTPUT_DOCX)
print(f"Report successfully generated at: {OUTPUT_DOCX}")
